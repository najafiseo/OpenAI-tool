import os
import openai
import streamlit as st
import pandas as pd
from deep_translator import GoogleTranslator
from bing_image_downloader import downloader
import gtts
from gtts import gTTS
from mutagen.mp3 import MP3
from PIL import Image
from pathlib import Path
import imageio
from moviepy import editor
import time
import urllib.request
import shutil
from docx import Document
import glob

api = st.text_input("")
openai.api_key = "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"
model_engine = "text-davinci-002"
option = st.selectbox(      # choosing 
    'choose one of the options:',
    ('How To Use The App' , 'summerize', 'Extract keywords', 'title generator', "generate article" ))
st.title("Enter ur API key")

if option == "Extract keywords":
    if api  :
        inp = st.text_input("Extract keywords from this text :")
        response = openai.Completion.create(
        model="text-davinci-003",
        prompt="Extract keywords from this text:\n\n" + inp,
        temperature=0.5,
        max_tokens=60,
        top_p=1.0,
        frequency_penalty=0.8,
        presence_penalty=0.0
        )
        for i in response["choices"]:
            key=i["text"][12:].split(",")
        dict ={"Keywords" : key}

        @st.cache
        def convert_df(df):
            return df.to_csv().encode('utf-8')

        my_large_df = pd.DataFrame(dict)
        csv = convert_df(my_large_df)
        st.download_button(
            label="Download data as CSV",
            data=csv,
            file_name='large_df.csv',
            mime='text/csv',
        )
        st.dataframe(my_large_df)
if option == "summerize":
    inpu = st.text_input("put ur text")
    inpu = GoogleTranslator(source="auto" , target= "en").translate(inpu)
    if inpu:
        response = openai.Completion.create(
        model="text-davinci-003",
        prompt="Summarize this for a second-grade student:\n\n"+ inpu,
        temperature=0.7,
        max_tokens=1000,
        top_p=1.0,
        frequency_penalty=0.0,
        presence_penalty=0.0
        )
        if response :
            re =response["choices"][0]["text"].replace("\n" , "")
            re_fa = GoogleTranslator(source="en" , target="fa").translate(re)
            st.success(re)
            st.warning(re_fa)
        else :
            st.warning("YOU HAVE PROBLEM SOMEWHERE")


if option == "generate article":
    nubmer = st.slider("how many paragragh do you need ? " , 0 ,10 )
    nubmer_questions = st.slider("how many questions do you need ?" , 0 , 10 )
    fa_or_en = st.selectbox("choose : " , ("persian" , "english"))
    AI_or_bing = st.selectbox("Generate image OR Download ready image" , ("download" , "generate"))
    if AI_or_bing == "generate":
        pass
    else :
        img_num = st.slider("how many image for every article do you need ? " , 0 , 99)
    uploaded_file = st.file_uploader("Choose a file")
    lisst = []
    if fa_or_en:
        if uploaded_file is not None:
            dataframe = pd.read_excel(uploaded_file)
            for index, row in dataframe.iterrows():
                if fa_or_en == "persian":
                    lisst.append(row['keyword'])
                    lisst = GoogleTranslator(source="auto" , target="en").translate_batch(lisst)
                else :
                    lisst.append(row['keyword'])
        for title_input in lisst :
            prompt ="give me " +  str(nubmer) +  " subheading about " + title_input
            completions = openai.Completion.create(
                engine=model_engine,
                prompt=prompt,
                max_tokens=1000,
                n=1,
                stop=None,
                temperature=0.7,
            )
            generated_text = completions.choices[0].text
            lines = generated_text.strip().split("\n")
            headers = [line for line in lines if line != '']
            headers_fa = GoogleTranslator(source="en" , target="fa").translate_batch(headers)
            li = []
            for i in headers :
                prompt = "write a paragraph about "+ i[3:]
                completions = openai.Completion.create(
                    engine=model_engine,
                    prompt=prompt,
                    max_tokens=1000,
                    n=1,
                    stop=None,
                    temperature=0.7,
                )      
                generated_paragraph = completions.choices[0].text
                pa = generated_paragraph.strip().split("\n")
                if pa != '' : 
                    li.append(pa)
            list_en = [x for xs in li for x in xs]
            list_fa_new=[]
            for items in list_en :
                fa_append = GoogleTranslator(source="en" , target="fa").translate(items)
                list_fa_new.append(fa_append)

            # making questions
            prompt = "discover the top " +  str(nubmer_questions) + " questions about " + title_input
            completions = openai.Completion.create(
                engine=model_engine,
                prompt=prompt,
                max_tokens=1000,
                n=1,
                stop=None,
                temperature=0.7,
            ) 
            generated_questions = completions.choices[0].text
            lines = generated_questions.strip().split("\n")
            questions = [line for line in lines if line != '']
            questions_fa = GoogleTranslator(source="en" , target="fa").translate_batch(questions)
            answer = []
            for question in questions :
                prompt = "write a paragraph about "+ question[3:]
                completions = openai.Completion.create(
                    engine="text-davinci-002",
                    prompt=prompt,
                    max_tokens=1000,
                    n=1,
                    stop=None,
                    temperature=0.7,
                )      
                generated_qu_paragraph = completions.choices[0].text
                qu_ans = generated_qu_paragraph.strip().split("\n")
                if qu_ans != '' : 
                    answer.append(qu_ans)

            list_en_answer = [x for xs in answer for x in xs]
            list_fa_answer_new=[]
            for items in list_en_answer :
                fa_append = GoogleTranslator(source="en" , target="fa").translate(items)
                list_fa_answer_new.append(fa_append)


            if fa_or_en == "persian":
                question_and_answer_dic = dict(zip(questions_fa , list_fa_answer_new))
                dic = dict(zip(headers_fa , list_fa_new))
            else :
                question_and_answer_dic = dict(zip(questions , list_en_answer))
                dic = dict(zip(headers , list_en))

            for key, value in dic.items():
                print(key[3:])
                print(value)
            for key, value in question_and_answer_dic.items():
                print(key)
                print(value)
                
                


            document = Document()       #make word
            for key, value in dic.items():
                document.add_heading(key[3:] , level =2)
                document.add_paragraph(value)
            for key, value in question_and_answer_dic.items():
                document.add_heading(key[3:] , level =2)
                document.add_paragraph(value)

            document_name = title_input + ".docx"
            document.save(document_name)
        # Download or Generating Image
        if AI_or_bing == "download" :
            for title_input in lisst :
                down = downloader.download(title_input, limit=img_num,  output_dir='dataset', adult_filter_off=True, force_replace=False, timeout=60, verbose=True)
        else:
            for title_input in lisst :
                li = [title_input , title_input , title_input]
                list =[]
                for i in li :                  
                    prompt = i
                    response = openai.Image.create(
                        prompt=prompt,
                        model="image-alpha-001"
                    )
                    list.append(response["data"][0]["url"])
                    file_name = ["img1" , "img2" , "img3" , "img4"]
                    file_path = "dataset/" + title_input +"\\"
                    if not os.path.exists(file_path):
                        os.makedirs(file_path)
                    def download_image(url, file_path ,file_name ):
                        full_path = file_path +file_name + ".jpg"
                        urllib.request.urlretrieve(url, full_path)               
                    for f_n , url in zip(file_name , list) :
                        download_image(url, file_path , f_n)


        for title_input in lisst:
            def text_from_docx(filename):
                doc = Document(filename)
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                return '\n'.join(fullText)
            document_name = title_input + ".docx"
            text = text_from_docx(document_name)
            if fa_or_en == "persian":
                tts = gTTS(text , lang = "ar")
            else :
                tts = gTTS(text , lang = "en")
            audio_name = title_input + ".mp3"
            tts.save(audio_name)



            audio_path = os.path.join(os.getcwd(), audio_name)
            video_path = os.path.join(os.getcwd(), "videos")
            path = "dataset\\" + title_input
            images_path = os.path.join(os.getcwd(), path)
            audio = MP3(audio_path)
            audio_length = audio.info.length

            list_of_images = []
            for image_file in os.listdir(images_path):
                if image_file.endswith('.png') or image_file.endswith('.jpg'):
                    image_path = os.path.join(images_path, image_file)
                    image = Image.open(image_path).resize((1200, 1200), Image.ANTIALIAS)
                    list_of_images.append(image)

            duration = audio_length/len(list_of_images)
            imageio.mimsave('images.gif', list_of_images, fps=1/duration)

            video = editor.VideoFileClip("images.gif")
            audio = editor.AudioFileClip(audio_path)
            final_video = video.set_audio(audio)
            filename = title_input + ".mp4"
            final_video.write_videofile(fps=60, codec="libx264", filename=filename)

            # cut to the path
            src = os.getcwd() + "\\" + document_name
            src_speech = os.getcwd() +r"\\" + audio_name
            src_vid = os.getcwd() + r"\\" + filename
            src_img = os.getcwd() + r"\dataset" + "\\" + title_input 
            
            if fa_or_en == "persian" :
                if not os.path.exists(title_input):
                    title_input = GoogleTranslator(source="auto" , target="fa").translate(title_input)
                    des = os.getcwd()+ r"\\" + title_input
                    os.mkdir(title_input)
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)
                    shutil.move(src_img , des)
                else : 
                    des = os.getcwd()+ r"\\" + title_input
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)
                    shutil.move(src_img , des)
            else :
                if not os.path.exists(title_input):
                    des = os.getcwd()+ r"\\" + title_input
                    os.mkdir(title_input)
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)
                    shutil.move(src_img , des)
                else : 
                    des = os.getcwd()+ r"\\" + title_input
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)   
                    shutil.move(src_img , des)           

            st.success ("Boooom your Article , Video & Audio is reeeeeaaaadyyyy :))))) ")

if option == "title generator":
    start_sequence = "\nAI:"
    restart_sequence = "\nHuman: "
    topic_input_first = st.text_input("give me your topic")
    topic_input = GoogleTranslator(source="auto" , target="en").translate(topic_input_first)
    if topic_input:
        generate_content = st.selectbox("do you want to generate content for the title" , ("choose one","yes" , "no"))
        if generate_content == "no" :
            openai.api_key = "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"
            response = openai.Completion.create(
            model="text-davinci-003",
            prompt="Human:give me a title about " + topic_input,
            temperature=0.9,
            max_tokens=150,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0.6,
            stop=[" Human:", " AI:"]
            )
            res =response ["choices"][0]["text"].replace("\n" , "")
            res_to_farsi = GoogleTranslator(source="auto" , target="fa").translate(res)
            res_to_en = GoogleTranslator(source="auto" , target="en").translate(res)
            st.success (res_to_en)
            st.warning(res_to_farsi)
        if generate_content == "yes":
            nubmer = st.slider("how many paragragh do you need ? " , 0 ,10 )
            nubmer_questions = st.slider("how many questions do you need ?" , 0 , 10 )

            AI_or_bing = st.selectbox("Generate image OR Download ready image" , ("download" , "generate"))
            if AI_or_bing == "generate":
                about_image = st.text_input("Tell me an image that you want : ")
            else :
                img_num = st.slider("how many image for every article do you need ? " , 0 , 99)
            
            fa_or_en = st.selectbox("choose : " , ( "choose" , "persian" , "english"))
            if fa_or_en == "choose":
                st.warning("you need to choose a lang")
            else :
                openai.api_key = "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"
                response = openai.Completion.create(
                model="text-davinci-003",
                prompt="Human:give me a title about " + topic_input,
                temperature=0.9,
                max_tokens=150,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0.6,
                stop=[" Human:", " AI:"]
                )
                res =response ["choices"][0]["text"].replace("\n" , "")
                res_to_farsi = GoogleTranslator(source="auto" , target="fa").translate(res)
                res_to_en = GoogleTranslator(source="auto" , target="en").translate(res)
                st.success(res_to_farsi)
                st.success(res_to_en)
                title_input = res_to_en

                prompt ="give me " +  str(nubmer) +  " subheading about " + title_input
                completions = openai.Completion.create(
                    engine=model_engine,
                    prompt=prompt,
                    max_tokens=1000,
                    n=1,
                    stop=None,
                    temperature=0.7,
                )
                generated_text = completions.choices[0].text
                lines = generated_text.strip().split("\n")
                headers = [line for line in lines if line != '']
                headers_fa = GoogleTranslator(source="en" , target="fa").translate_batch(headers)
                li = []
                for i in headers :
                    prompt = "write a paragraph about "+ i[3:]
                    completions = openai.Completion.create(
                        engine=model_engine,
                        prompt=prompt,
                        max_tokens=1000,
                        n=1,
                        stop=None,
                        temperature=0.7,
                    )      
                    generated_paragraph = completions.choices[0].text
                    pa = generated_paragraph.strip().split("\n")
                    if pa != '' : 
                        li.append(pa)
                list_en = [x for xs in li for x in xs]
                list_fa_new=[]
                for items in list_en :
                    fa_append = GoogleTranslator(source="en" , target="fa").translate(items)
                    list_fa_new.append(fa_append)

                # making questions
                prompt = "discover the top " +  str(nubmer_questions) + " questions about " + title_input
                completions = openai.Completion.create(
                    engine=model_engine,
                    prompt=prompt,
                    max_tokens=1000,
                    n=1,
                    stop=None,
                    temperature=0.7,
                ) 
                generated_questions = completions.choices[0].text
                lines = generated_questions.strip().split("\n")
                questions = [line for line in lines if line != '']
                questions_fa = GoogleTranslator(source="en" , target="fa").translate_batch(questions)
                answer = []
                for question in questions :
                    prompt = "write a paragraph about "+ question[3:]
                    completions = openai.Completion.create(
                        engine="text-davinci-002",
                        prompt=prompt,
                        max_tokens=1000,
                        n=1,
                        stop=None,
                        temperature=0.7,
                    )      
                    generated_qu_paragraph = completions.choices[0].text
                    qu_ans = generated_qu_paragraph.strip().split("\n")
                    if qu_ans != '' : 
                        answer.append(qu_ans)

                list_en_answer = [x for xs in answer for x in xs]
                list_fa_answer_new=[]
                for items in list_en_answer :
                    fa_append = GoogleTranslator(source="en" , target="fa").translate(items)
                    list_fa_answer_new.append(fa_append)


                if fa_or_en == "persian":
                    question_and_answer_dic = dict(zip(questions_fa , list_fa_answer_new))
                    dic = dict(zip(headers_fa , list_fa_new))
                else :
                    question_and_answer_dic = dict(zip(questions , list_en_answer))
                    dic = dict(zip(headers , list_en))
                    
                    
                document = Document()       #make word
                document.add_heading(title_input , level = 1)
                for key, value in dic.items():
                    document.add_heading(key[3:] , level =2)
                    document.add_paragraph(value)
                for key, value in question_and_answer_dic.items():
                    document.add_heading(key[3:] , level =2)
                    document.add_paragraph(value)

                document_name = topic_input_first + ".docx"
                document.save(document_name)

                # Download or Generating Image
                if AI_or_bing == "download" :
                        down = downloader.download(topic_input_first, limit=img_num,  output_dir='dataset', adult_filter_off=True, force_replace=False, timeout=60, verbose=True)
                else:
                    li = [title_input , title_input , about_image]
                    list =[]
                    for i in li :                  
                        prompt = i
                        response = openai.Image.create(
                            prompt=prompt,
                            model="image-alpha-001"
                        )
                        list.append(response["data"][0]["url"])
                        file_name = ["img1" , "img2" , "img3" , "img4"]
                        file_path = "dataset/" + topic_input_first +"\\"
                        if not os.path.exists(file_path):
                            os.makedirs(file_path)
                        def download_image(url, file_path ,file_name ):
                            full_path = file_path +file_name + ".jpg"
                            urllib.request.urlretrieve(url, full_path)               
                        for f_n , url in zip(file_name , list) :
                            download_image(url, file_path , f_n)


                # voice
                def text_from_docx(filename):
                    doc = Document(filename)
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    return '\n'.join(fullText)
                document_name = topic_input_first + ".docx"
                text = text_from_docx(document_name)
                if fa_or_en == "persian":
                    tts = gTTS(text , lang = "ar")
                else :
                    tts = gTTS(text , lang = "en")
                audio_name = topic_input_first + ".mp3"
                tts.save(audio_name)



                audio_path = os.path.join(os.getcwd(), audio_name)
                video_path = os.path.join(os.getcwd(), "videos")
                path = "dataset\\" + topic_input_first
                images_path = os.path.join(os.getcwd(), path)
                audio = MP3(audio_path)
                audio_length = audio.info.length

                list_of_images = []
                for image_file in os.listdir(images_path):
                    if image_file.endswith('.png') or image_file.endswith('.jpg'):
                        image_path = os.path.join(images_path, image_file)
                        image = Image.open(image_path).resize((1200, 1200), Image.ANTIALIAS)
                        list_of_images.append(image)

                duration = audio_length/len(list_of_images)
                imageio.mimsave('images.gif', list_of_images, fps=1/duration)

                video = editor.VideoFileClip("images.gif")
                audio = editor.AudioFileClip(audio_path)
                final_video = video.set_audio(audio)
                filename = topic_input_first + ".mp4"
                final_video.write_videofile(fps=60, codec="libx264", filename=filename)

                # cut to the path
                src = os.getcwd() + "\\" + document_name
                src_speech = os.getcwd() +r"\\" + audio_name
                src_vid = os.getcwd() + r"\\" + filename
                src_img = os.getcwd() + r"\dataset" + "\\" + topic_input_first 
                des = os.getcwd()+ "\\" + topic_input_first 
                if not os.path.exists(topic_input_first):
                    os.mkdir(topic_input_first)
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)
                    shutil.move(src_img , des)
                else : 
                    shutil.move(src , des)
                    shutil.move(src_speech , des)
                    shutil.move(src_vid , des)                

                st.success ("Boooom your Article , Video & Audio is reeeeeaaaadyyyy :))))) ")
